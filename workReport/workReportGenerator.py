# -*- coding: utf-8 -*-
import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK


# valign: "top","center","bottom"
def setCellStyle(cell, text, fontSize, isBold, isItalic, isUnderlined, valign="center",
                 align=WD_PARAGRAPH_ALIGNMENT.CENTER):
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(fontSize)
    run.bold = isBold
    run.italic = isItalic
    run.underlined = isUnderlined
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), valign)
    tcPr.append(tcVAlign)


def generateDocument():
    # создаём документ
    document = Document()
    # задаём отступы
    document.sections[0].top_margin = 200000
    document.sections[0].left_margin = 650000
    document.sections[0].right_margin = 200000

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'

    return document


# генерировать первую страницу документа
# supervisorShort - инициалы руководителя
# workerNumber - табельный номер работника
# workerPosition - должность работника
# rationales - обоснования (номера, обоснования)
# works - описание работ (номер, участок, выполняемые работы, начало, конец)
# factWorks - описание фактически выполненных работ (номер, участок, выполняемые работы, начало, конец)
# reportMaker - составитель отчёта
# reportChecker - принял отчёт
# reportVIKer - ВИК
# attestation - аттестация
def generateWorkReportPage1(document, supervisorShort, workerName, workerNumber, date, workerPosition,
                            rationales, works, factWorks, reportMaker, reportChecker,
                            reportVIKer, note, attestation):
    # создаём таблицу
    table = document.add_table(rows=4, cols=0, style=document.styles['TableGrid'])

    table.add_column(width=Inches(0.9))
    table.add_column(width=Inches(1.1))
    table.add_column(width=Inches(1.2))
    table.add_column(width=Inches(3.7))
    table.add_column(width=Inches(1))
    table.add_column(width=Inches(0.5))

    cell = table.cell(0, 0).merge(table.cell(0, 1)).merge(table.cell(0, 2)).merge(table.cell(0, 3))
    setCellStyle(cell, 'Сменный наряд ' + supervisorShort + '-' + workerNumber + '-' + str(date), 20, True, False,
                 False)
    cell = table.cell(0, 4).merge(table.cell(0, 5))
    setCellStyle(cell, '1/2', 19 + len(rationales), True, False, False)

    cell = table.cell(1, 0).merge(table.cell(1, 1))
    setCellStyle(cell, 'Должность', 10, True, False, False)

    cell = table.cell(2, 0).merge(table.cell(2, 1))
    setCellStyle(cell, 'С нарядом ознакомился:', 10, True, False, False)

    setCellStyle(table.cell(1, 2), workerPosition, 10, False, False, False)
    cell = table.cell(1, 3).merge(table.cell(2, 3))
    setCellStyle(table.cell(1, 3), workerName, 16, False, False, False)
    table.cell(1, 4).merge(table.cell(1, 5))
    table.cell(2, 4).merge(table.cell(2, 5))
    setCellStyle(table.cell(1, 4), 'Сведения об аттестации', 10, False, False, False)
    setCellStyle(table.cell(2, 4), attestation, 10, False, False, False)

    setCellStyle(table.cell(3, 0), '№п/п', 10, True, False, False)
    table.cell(3, 1).merge(table.cell(3, 5))
    setCellStyle(table.cell(3, 1), 'Обоснование для выполнения работ:', 10, True, False, False)

    pos = 3
    if len(rationales) == 0:
        table.add_row()
        pos += 1
        setCellStyle(table.cell(pos, 0), "", 10, False, False, False)
        table.cell(pos, 1).merge(table.cell(pos, 5))
        setCellStyle(table.cell(pos, 1), "", 10, False, False, False)

    for rationale in rationales:
        table.add_row()
        pos += 1
        setCellStyle(table.cell(pos, 0), rationale[0], 10, False, False, False)
        table.cell(pos, 1).merge(table.cell(pos, 5))
        setCellStyle(table.cell(pos, 1), rationale[1], 10, False, False, False)

    pos += 1
    table.add_row()
    setCellStyle(table.cell(pos, 0), '№п/п', 10, True, False, False)
    setCellStyle(table.cell(pos, 1), 'Раб. место', 10, True, False, False)
    table.cell(pos, 2).merge(table.cell(pos, 3))
    setCellStyle(table.cell(pos, 2), 'Выполняемые работы', 10, True, False, False)
    table.cell(pos, 4).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 4), 'Период', 10, True, False, False)

    pos += 1

    for work in works:
        table.add_row()
        table.add_row()
        for i in range(4):
            table.cell(pos, i).merge(table.cell(pos + 1, i))
        setCellStyle(table.cell(pos, 0), work[0], 10, False, False, False)
        setCellStyle(table.cell(pos, 1), work[1], 10, False, False, False)
        table.cell(pos, 2).merge(table.cell(pos, 3))
        setCellStyle(table.cell(pos, 2), work[2], 10, False, False, False)
        setCellStyle(table.cell(pos, 4), 'Начало', 10, True, False, False)
        setCellStyle(table.cell(pos + 1, 4), 'Конец', 10, True, False, False)
        setCellStyle(table.cell(pos, 5), work[3], 10, False, False, False)
        setCellStyle(table.cell(pos + 1, 5), work[4], 10, False, False, False)
        pos += 2

    table.add_row()
    setCellStyle(table.cell(pos, 0), '№п/п', 10, True, False, False)
    setCellStyle(table.cell(pos, 1), 'Раб. место', 10, True, False, False)
    table.cell(pos, 2).merge(table.cell(pos, 3))
    setCellStyle(table.cell(pos, 2), 'Фактически выполненные работы, причина невыполнения работ', 10, True, False,
                 False)
    table.cell(pos, 4).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 4), 'Период(факт.)', 10, True, False, False)

    pos += 1
    for work in factWorks:
        table.add_row()
        table.add_row()
        for i in range(4):
            table.cell(pos, i).merge(table.cell(pos + 1, i))
        setCellStyle(table.cell(pos, 0), work[0], 10, False, False, False)
        setCellStyle(table.cell(pos, 1), work[1], 10, False, False, False)
        table.cell(pos, 2).merge(table.cell(pos, 3))
        setCellStyle(table.cell(pos, 2), work[2], 10, False, False, False)
        setCellStyle(table.cell(pos, 4), 'Начало', 10, True, False, False)
        setCellStyle(table.cell(pos + 1, 4), 'Конец', 10, True, False, False)
        setCellStyle(table.cell(pos, 5), work[3], 10, False, False, False)
        setCellStyle(table.cell(pos + 1, 5), work[4], 10, False, False, False)
        pos += 2

    table.add_row()
    table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(table.cell(pos, 0), 'Сменный наряд составил:', 12, False, False, False)
    table.cell(pos, 4).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 4), reportMaker, 12, True, False, False)

    pos += 1
    table.add_row()
    table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(table.cell(pos, 0), 'Работы по наряду принял:', 12, False, False, False)
    table.cell(pos, 4).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 4), reportChecker, 12, True, False, False)

    pos += 1
    table.add_row()
    table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(table.cell(pos, 0), 'ВИК (при необходимости):', 12, False, False, False)
    table.cell(pos, 4).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 4), reportVIKer, 12, True, False, False)

    pos += 1
    table.add_row()
    table.cell(pos, 0).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 0), 'Примечания', 10, True, False, False)

    pos += 1
    table.add_row()
    table.cell(pos, 0).merge(table.cell(pos, 5))
    setCellStyle(table.cell(pos, 0), note, 16, True, False, False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    table.cell(pos, 0).add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    table.columns[0].width = 350


# генерировать вторую страницу документа
# supervisorShort - инициалы руководителя
# workerNumber - табельный номер работника
# date  - дата
# planEquipment - плановая выдача оборудования
# nonPlanEquipment - внеплановая выдача оборудования
# dust - направлено в изолятор брака
# supervisorName - имя руководителя
# stockManName - имя кладовщика
def generateWorkReportPage2(document, supervisorShort, workerName, workerNumber, date, planEquipment, nonPlanEquipment,
                            dust, supervisorName, stockManName):
    # создаём таблицу
    table = document.add_table(rows=5, cols=0, style=document.styles['TableGrid'])

    table.add_column(width=Inches(0.4))
    table.add_column(width=Inches(1.5))
    table.add_column(width=Inches(1.5))
    table.add_column(width=Inches(0.5))
    table.add_column(width=Inches(0.8))
    table.add_column(width=Inches(0.8))
    table.add_column(width=Inches(0.8))
    table.add_column(width=Inches(0.8))
    table.add_column(width=Inches(0.8))

    cell = table.cell(0, 0).merge(table.cell(0, 6))
    setCellStyle(cell, 'Сменный наряд ' + supervisorShort + '-' + workerNumber + '-' + str(date), 20, True, False,
                 False)

    cell = table.cell(0, 7).merge(table.cell(0, 8))
    setCellStyle(cell, '2/2', 20, True, False,
                 False)

    cell = table.cell(1, 0).merge(table.cell(1, 8))
    setCellStyle(cell, workerName, 16, False, False,
                 False)

    cell = table.cell(2, 0).merge(table.cell(2, 8))
    setCellStyle(cell, 'Выдача комплектующих и оснастки плановая', 11, True, False,
                 False)

    cell = table.cell(3, 0).merge(table.cell(4, 0))
    setCellStyle(cell, '№п/п', 10, True, False,
                 False)

    cell = table.cell(3, 1).merge(table.cell(4, 1))
    setCellStyle(cell, 'Наименование', 10, True, False,
                 False)

    cell = table.cell(3, 2).merge(table.cell(4, 2))
    setCellStyle(cell, 'Шифр', 10, True, False,
                 False)

    cell = table.cell(3, 3).merge(table.cell(3, 4))
    setCellStyle(cell, 'Передано для выполнения работ', 10, True, False,
                 False)

    setCellStyle(table.cell(4, 3), 'Ед. изм.', 10, True, False,
                 False)

    setCellStyle(table.cell(4, 4), 'Кол-во', 10, True, False,
                 False)

    cell = table.cell(3, 5).merge(table.cell(4, 5))
    setCellStyle(cell, 'Использовано', 10, True, False,
                 False)

    cell = table.cell(3, 6).merge(table.cell(4, 6))
    setCellStyle(cell, 'Брак', 10, True, False,
                 False)

    cell = table.cell(3, 7).merge(table.cell(3, 8))
    setCellStyle(cell, 'Передано на склад', 10, True, False,
                 False)

    setCellStyle(table.cell(4, 7), 'Утиль', 10, True, False,
                 False)

    setCellStyle(table.cell(4, 8), 'Остаток', 10, True, False,
                 False)
    pos = 4
    for pl in planEquipment:
        pos += 1
        table.add_row()
        setCellStyle(table.cell(pos, 0), str(pos - 4), 10, False, False, False)
        for i in range(1, 9):
            setCellStyle(table.cell(pos, i), pl[i - 1], 10, False, False, False)

    pos += 1
    table.add_row()

    cell = table.cell(pos, 0).merge(table.cell(pos, 8))
    setCellStyle(cell, 'Выдача комплектующих и оснастки внеплановая', 11, True, False,
                 False)
    j = 1
    for pl in nonPlanEquipment:
        pos += 1
        table.add_row()
        setCellStyle(table.cell(pos, 0), str(j), 10, False, False, False)
        j += 1
        for i in range(1, 9):
            setCellStyle(table.cell(pos, i), pl[i - 1], 10, False, False, False)

    pos += 1
    table.add_row()
    cell = table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(cell, supervisorName, 10, True, False, False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    setCellStyle(table.cell(pos, 3), 'Утв.', 10, True, False, False)
    setCellStyle(table.cell(pos, 5), 'Утв.', 10, True, False, False)

    pos += 1
    table.add_row()
    cell = table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(cell, workerName, 10, True, False, False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    setCellStyle(table.cell(pos, 3), 'Согл.', 10, True, False, False)
    setCellStyle(table.cell(pos, 5), 'Согл.', 10, True, False, False)

    table.cell(pos - 1, 7).merge(table.cell(pos, 8))

    pos += 1
    table.add_row()
    cell = table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(cell, workerName, 10, True, False, False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    setCellStyle(table.cell(pos, 3), 'Получ.', 10, True, False, False)
    setCellStyle(table.cell(pos, 7), 'Получ.', 10, True, False, False)

    pos += 1
    table.add_row()
    cell = table.cell(pos, 0).merge(table.cell(pos, 2))
    setCellStyle(cell, stockManName, 10, True, False, False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    setCellStyle(table.cell(pos, 3), 'Выдал', 10, True, False, False)
    setCellStyle(table.cell(pos, 7), 'Перед.', 10, True, False, False)

    table.cell(pos - 1, 5).merge(table.cell(pos, 6))

    pos += 1
    table.add_row()
    cell = table.cell(pos, 0).merge(table.cell(pos, 8))
    setCellStyle(cell, 'Направлено в изолятор брака', 11, True, False,
                 False)

    pos += 1
    table.add_row()
    cell = table.cell(pos, 1).merge(table.cell(pos, 6))
    setCellStyle(cell, 'Наименование', 11, True, False,
                 False)
    cell = table.cell(pos, 7).merge(table.cell(pos, 8))
    setCellStyle(cell, 'Кол-во', 11, True, False,
                 False)

    setCellStyle(table.cell(pos, 0), '№п/п', 10, True, False,
                 False)
    j = 1
    for pl in dust:
        pos += 1
        table.add_row()
        setCellStyle(table.cell(pos, 0), str(j), 10, False, False, False)
        j += 1
        setCellStyle(table.cell(pos, 1).merge(table.cell(pos, 6)), pl[0], 10, False, False, False)
        setCellStyle(table.cell(pos, 7).merge(table.cell(pos, 8)), pl[1], 10, False, False, False)

    pos += 1
    table.add_row()
    setCellStyle(table.cell(pos, 1), 'Передал:', 11, True, False,
                 False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    setCellStyle(table.cell(pos, 3).merge(table.cell(pos, 4)), supervisorName, 11, True, False,
                 False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    pos += 1
    table.add_row()
    setCellStyle(table.cell(pos, 1), 'Принял:', 11, True, False,
                 False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    setCellStyle(table.cell(pos, 3).merge(table.cell(pos, 4)), stockManName, 11, True, False,
                 False, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    table.cell(pos - 1, 5).merge(table.cell(pos, 8))


def generateReport(supervisorShort, supervisorName, workerName, workerNumber, stockManName, date, workerPosition,
                   rationales, works, factWorks, reportMaker, reportChecker,
                   reportVIKer, note, attestation, dust, planEquipment, nonPlanEquipment):
   # print(workerPosition)
    #print(attestation)
    document = generateDocument()
    generateWorkReportPage1(document, supervisorShort, workerName, workerNumber, date, workerPosition,
                            rationales, works, factWorks, reportMaker, reportChecker,
                            reportVIKer, note, attestation)
    document.add_page_break()
    generateWorkReportPage2(document, supervisorShort, workerName, workerNumber, date, planEquipment, nonPlanEquipment,
                            dust, supervisorName, stockManName)
    return document
